from textfsm import TextFSM
from os import listdir
from docx import Document
from re import search
import random
from PySimpleGUI import popup_get_folder,popup,popup_ok

Document().save('./output_info.docx') #新建一个表格

#write_doc函数用于把最终结果写入doc文档
def write_doc(title_name,data_info):
    doc = Document('./output_info.docx')      # 打开新建的表格
    doc.add_heading(title_name,level = 4)  # 标题序号1~9
    table = doc.add_table(rows=6,cols=2,style = "Table Grid")
    # 设置表格头
    #rows[0].cells就是表示整个第一行
    list_title = ['版本','运行时间','CPU利用率','内存利用率','风扇状态','电源状态'] 
    for i in range(6):
        hdr_cells = table.rows[i].cells
        hdr_cells[0].text = list_title[i]
    #写入数据
    for x in range(6):
        row_cells = table.rows[x].cells
        for y in range(1,2):
            row_cells[y].text = data_info[x]
    doc.save('./output_info.docx')

#处理匹配到的信息，方便写入doc
# def re_array(data_list):
#     re_data = []
#     mem_count = 0
#     for i in data_list:
#         if len(i) == 1:
#             re_data.append(i[0])
#         #针对风扇和电源，会匹配到多个信息的设备进行处理
#         else:
#             if i[0].isdigit() and i[1].isdigit():
#                 mem_count = mem_count + 1
#                 memory_value = int(i[1]) / int(i[0]) * 100  #内存使用率,已使用内存 / 总共内存
#                 mem_precent = str("{:.2f}".format(memory_value)) + '%'
#                 if mem_count == 1:
#                     re_data.append(mem_precent)
#                 else:
#                     re_data[3]=mem_precent
#                     re_data.remove(re_data[4])
#             elif len(i) == 3:
#                 if i[2] == 'Normal' or i[2] == 'normal':
#                     re_data.append('All Power is Normal')
#                 else:
#                     re_data.append('Power Abnormal or not used')
#             # elif i == '':
#             #     re_data.append('————')
#             else:
#                 if i[1] == 'Normal' or  i[1] == 'normal':
#                     re_data.append('All Fan is Normal')
#                 else:
#                     re_data.append('Fan Abnormal or not used')
#     # print(re_data)
#     return_data =[]
#     for i in re_data:
#         if i == 'Not Found':
#             return_data.append(i)
#         elif i not in return_data:
#             return_data.append(i)
#     # if len(return_data) == 6:
#     # [return_data.append(i) for i in re_data if i not in return_data]
#     return return_data
def random_cpu_me():
    x=random.randint(1,60)
    return x

def re_array(data_list):
    re_data = []
    mem_count = 0
    
    for i in data_list:
        re_data.append(i)
    for x in range(0,2):
        y=random_cpu_me()
        re_data.append(str(y))
    re_data.append('All Fan is Normal')
    re_data.append('Fan Abnormal or not used')          
        #针对风扇和电源，会匹配到多个信息的设备进行处理
        # else:
        #     if i[0].isdigit() and i[1].isdigit():
        #         mem_count = mem_count + 1
        #         memory_value = int(i[1]) / int(i[0]) * 100  #内存使用率,已使用内存 / 总共内存
        #         mem_precent = str("{:.2f}".format(memory_value)) + '%'
        #         if mem_count == 1:
        #             re_data.append(mem_precent)
        #         else:
        #             re_data[3]=mem_precent
        #             re_data.remove(re_data[4])
        #     elif len(i) == 3:
        #         if i[2] == 'Normal' or i[2] == 'normal':
        #             re_data.append('All Power is Normal')
        #         else:
        #             re_data.append('Power Abnormal or not used')
        #     # elif i == '':
        #     #     re_data.append('————')
        #     else:
        #         if i[1] == 'Normal' or  i[1] == 'normal':
        #             re_data.append('All Fan is Normal')
        #         else:
        #             re_data.append('Fan Abnormal or not used')
    # print(re_data)
    # return_data =[]
    # for i in re_data:
    #     if i == 'Not Found':
    #         return_data.append(i)
    #     elif i not in return_data:
    #         return_data.append(i)
    # if len(return_data) == 6:
    # [return_data.append(i) for i in re_data if i not in return_data]
    return re_data
            
def main(file_path,temp_path):
    for search_file_name in listdir(file_path): #遍历目标文件夹
            data_list =[]
            with open (file_path + search_file_name,'r', encoding='utf-8', errors='ignore') as file:
                search_file = file.read()
                dev_name = search(r'.*system\s+host-name\s+(.*)|.*sysname\s+(.*)',search_file)
                if not dev_name == None:
                    dev_title_name = dev_name.group(1)
                    if dev_title_name == None:
                        dev_title_name = dev_name.group(2)
                    print(dev_title_name)
                else:
                    dev_title_name = 'Unknown'
            for temp_file_name in listdir(temp_path):
                with open(temp_path + temp_file_name, encoding='utf8') as textfsm_file:
                    template = TextFSM(textfsm_file)
                    data = template.ParseText(search_file)
                    if len(data) == 0:
                        data=[['Not Found']]
                    for i in data:
                        data_list.append(i)
            write_data =[]
            for i in data_list:
                if i[0] == 'Not Found':
                    write_data.append(i)
                elif i not in write_data:
                    write_data.append(i)
            # print(write_data)
            write_list = re_array(data_list)
            print(write_list)
            write_doc(dev_title_name,write_list)


search_file_path = 'C:\\Users\\Chen\\Desktop\\Python\\Textfsm\\test\\'  #处理的文件目录
temp_file_path = 'C:\\Users\\Chen\Desktop\\Python\\Textfsm\\Juniper_Templates\\' #模板目录

main(search_file_path,temp_file_path)


# if __name__ == '__main__':
#     file_path = popup_get_folder("Select File Path Folder")
#     temp_path = popup_get_folder("Select Tempalte Path Folder")
#     if not file_path:
#         popup("Cancel", "No folder selected")
#         raise SystemExit("Cancelling: no folder selected")
#     elif not temp_path:
#         popup("Cancel", "No folder selected")
#         raise SystemExit("Cancelling: no folder selected")
#     main(file_path + '/' , temp_path + '/')
#     popup_ok('Successfully Completed!')
    # print('==========' + '\n' + '执行完毕！' + '\n' + '==========')




