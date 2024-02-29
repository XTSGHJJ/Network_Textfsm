from textfsm import TextFSM
from docx import Document
# from docx.document import Document
from re import search,IGNORECASE,findall
from os import getcwd,listdir
from PySimpleGUI import popup_get_folder,popup,popup_ok

doc = Document('./template/ABC_temp.docx') #必须要有模板
paragraphs = doc.paragraphs

#移动表格到指定文本位置,原始代码理解学习
#_tbl 和 _p 方法是操作底层xml，获取到表格对象和段落对象对应的 XML 元素，用addnext 方法将表格添加到段落的下一个位置
# def move_table_after(table, paragraph):  #table是表格的内存地址，paragr是段落内存地址
#     tbl, p = table._tbl, paragraph._p 
#     p.addnext(tbl)

def move_element_after(element, paragraph): #element是表格或者标题的内存地址，paragr是段落内存地址
    if hasattr(element, '_tbl'):   #移动表格
        paragraph._p.addnext(element._tbl)
    elif hasattr(element, '_element'): #移动标题
        paragraph._p.addnext(element._element)

#创建设备硬件状态表
def device_status_table (file_numb,dev_name,table_title,find_tag):
    print(find_tag,table_title)
    print(type(find_tag),type(table_title))
    for doc_par in doc.paragraphs:
        if findall(find_tag,doc_par.text):
            dst_table= doc.add_table(rows=file_numb+1,cols=4,style = "Table Grid")
            # table_title = ['设备名','主引擎指示灯状态','电源指示灯状态','风扇指示灯状态']
            for table_num in range(file_numb+1):
                if table_num == 0:
                    title_hdr = dst_table.rows[table_num].cells
                    for title_cont in range(4):
                        title_hdr[title_cont].text = table_title[title_cont]
                else:
                    table_hdr = dst_table.rows[table_num].cells
                    for table_cont in range(4):
                        if table_cont == 0:
                            table_hdr[table_cont].text = dev_name[table_num-1]
                        else:
                            table_hdr[table_cont].text = '正常'
            move_element_after(dst_table,doc_par)
            break

#每个设备生成一个表格
def dev_run_info_tabl(title_name,data_info):
    title=doc.add_heading(title_name,level = 4)  # 标题序号1~9
    table = doc.add_table(rows=6,cols=2,style = "Table Grid") #添加表格
    # 设置表格头
    #rows[0].cells就是表示整个第一行
    list_title = ['版本','运行时间','CPU利用率','内存利用率','风扇状态','电源状态'] 
    for i in range(6):
        hdr_cells = table.rows[i].cells
        hdr_cells[0].text = list_title[i]
    #写入数据
    for x in range(6):
        row_cells = table.rows[x].cells
        for y in range(1,2):
            row_cells[y].text = data_info[x]
    tit_tab_info=[table,title]
    for doc_par in doc.paragraphs:
        if findall('标识1',doc_par.text):
            for move_content in tit_tab_info:
                move_element_after(move_content,doc_par)
            
#处理匹配到的信息，方便写入doc
def re_array(data_list):
    re_data = []
    mem_count = 0
    for i in data_list:
        if len(i) == 1:
            re_data.append(i[0])
        #针对风扇和电源，会匹配到多个信息的设备进行处理
        else:
            #临时处理运行时间和型号部分
            if len(i) == 2 and i[0].isdigit():
                re_data.append(i[0])


            if i[0].isdigit() and i[1].isdigit():
                mem_count = mem_count + 1
                memory_value = int(i[1]) / int(i[0]) * 100  #内存使用率,已使用内存 / 总共内存
                mem_precent = str("{:.2f}".format(memory_value)) + '%'
                if mem_count == 1:
                    re_data.append(mem_precent)
                else:
                    re_data[3]=mem_precent
                    re_data.remove(re_data[4])
            elif len(i) == 3:
                if i[2] == 'Normal' or i[2] == 'normal':
                    re_data.append('All Power is Normal')
                else:
                    re_data.append('Power Abnormal or not used')
            # elif i == '':
            #     re_data.append('————')
            else:
                if i[1] == 'Normal' or  i[1] == 'normal':
                    re_data.append('All Fan is Normal')
                else:
                    re_data.append('Fan Abnormal or not used')
    # print(re_data)
    return_data = []
    for i in re_data:
        if i == 'Not Found':
            return_data.append(i)
        elif i not in return_data:
            return_data.append(i)
    # if len(return_data) == 6:
    # [return_data.append(i) for i in re_data if i not in return_data]
    return return_data

def judge(txt):
    pwd_dir=getcwd().replace('\\','/')
    dev_type = {'Huawei\s+Technologies':'1','H3C\s+Comware':'2','Cisco':'3','JUNOS\s+Software':'4'}
    for  type in dev_type.keys():
        a=search(type,txt,IGNORECASE)
        if not a == None: 
            if dev_type[type] == '3':
                devname = search(r'.*hostname\s+(.*)',txt)
                if devname == None:
                    devname = search(r'(.*)#',txt)
                tem_path = pwd_dir + '/Cisco_Templates/'
                return devname,tem_path
            elif dev_type[type] == '4':
                devname = search(r'.*system\s+host-name\s+(.*)',txt)
                tem_path = pwd_dir + '/Juniper_Templates/'
                return devname,tem_path
            else:
                devname = search(r'sysname\s+(.*)',txt)
                tem_path = pwd_dir + '/H3c_Templates/'
                return devname,tem_path
    return 'Unknown' 

def main(file_path):
    file_num = len(listdir(file_path))
    dev_name_list = [] #设备名列表
    for search_file_name in listdir(file_path): #遍历目标文件夹
        data_list = [] #获取的内容
        with open (file_path + search_file_name,'r', encoding='utf-8', errors='ignore') as file_text:
            search_file = file_text.read()
            judge_val = judge(search_file)
            if judge_val == 'Unknown':
                continue            
            dev_title_name = judge_val[0].group(1) #提取匹配到信息设备名
            dev_name_list.append(dev_title_name)
            temp_path = judge_val[1]
        for temp_file_name in listdir(temp_path):
            with open(temp_path + temp_file_name, encoding='utf8') as textfsm_file:
                template = TextFSM(textfsm_file)
                data = template.ParseText(search_file) #使用模板获取文本信息

                if len(data) == 0:
                    data=[['Not Found']]
                for i in data:
                    data_list.append(i)
        write_list = re_array(data_list) #整合数据
        dev_run_info_tabl(dev_title_name,write_list) #生成表格信息

    dev_hd_table_title = ['设备名','主引擎指示灯状态','电源指示灯状态','风扇指示灯状态']
    dev_run_table_title = ['设备名','电源运行状态','风扇运行状态','Log日志分析']
    dev_tab_tit = {'表格1':dev_hd_table_title,'表格2':dev_run_table_title}
    for tag,tab_tit in dev_tab_tit.items():
        print(tag,tab_tit)
        device_status_table(file_num,dev_name_list,tab_tit,tag)
    doc.save('./ABC_output.docx')

# search_file_path = 'C:/Users/chen/Desktop/Python/Network_Textfsm/testfile/' #模板目录
# main(search_file_path)

if __name__ == '__main__':
    file_path = popup_get_folder("Select File Path Folder")
    # temp_path = popup_get_folder("Select Tempalte Path Folder")
    if not file_path:
        popup("Cancel", "No folder selected")
        raise SystemExit("Cancelling: no folder selected")
    # elif not temp_path:
    #     popup("Cancel", "No folder selected")
    #     raise SystemExit("Cancelling: no folder selected")
    main(file_path + '/')
    popup_ok('Successfully Completed!')
    print('==========' + '\n' + '执行完毕！' + '\n' + '==========')

